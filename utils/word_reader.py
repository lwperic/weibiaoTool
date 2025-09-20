"""
通用的 Word 文档读取工具

功能：
- 读取 .docx：使用 python-docx 提取表格与段落内容
- 读取 .doc：在 Windows 上优先通过 Word COM 自动转换为 .docx 再读取；
           若不可用则抛出明确异常，提示用户转换后再上传
"""

from __future__ import annotations

import os
import uuid
from pathlib import Path
from typing import List


def _read_docx_text(docx_path: Path) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(docx_path)
    content: List[str] = []

    # 表格
    for table in doc.tables:
        for row in table.rows:
            row_text: List[str] = []
            for cell in row.cells:
                row_text.append((cell.text or "").strip())
            if any(t for t in row_text):
                content.append("\t".join(row_text))

    # 段落
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            content.append(text)

    return "\n".join(content)


def _is_windows() -> bool:
    return os.name == "nt"


def _convert_doc_to_docx_via_com(doc_path: Path) -> Path:
    """使用 Windows Word COM 将 .doc 转换为 .docx。

    返回转换后的 .docx 路径。
    """
    if not _is_windows():
        raise RuntimeError("当前操作系统不支持 .doc 自动转换，请手动转换为 .docx 后再试。")

    try:
        import pythoncom
        from win32com.client import gencache, constants
    except Exception as exc:  # pragma: no cover - 环境相关
        raise RuntimeError(
            "未安装 pywin32 或无法导入 win32com，无法自动转换 .doc。" \
            "请安装 pywin32 或将文档手动转换为 .docx 后再上传。"
        ) from exc

    # 初始化COM
    pythoncom.CoInitialize()
    
    # 目标路径放在同目录下，防止权限问题，使用随机文件名避免冲突
    dest_path = doc_path.with_suffix("")
    dest_path = dest_path.parent / f"{dest_path.name}.{uuid.uuid4().hex}.converted.docx"

    word = None
    doc = None
    try:
        word = gencache.EnsureDispatch("Word.Application")
        word.Visible = False
        # 读/写路径必须是绝对路径且使用 str
        doc = word.Documents.Open(str(doc_path.absolute()))
        # 16: wdFormatXMLDocument
        doc.SaveAs(str(dest_path.absolute()), FileFormat=16)
        doc.Close(False)
        doc = None
        return dest_path
    finally:
        try:
            if doc is not None:
                doc.Close(False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        # 释放COM
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


def read_word_text(file_path: Path) -> str:
    """读取 Word 文本内容，自动兼容 .doc 与 .docx。

    - .docx：直接解析
    - .doc：在 Windows 上尝试通过 COM 转换为 .docx 后解析
    """
    suffix = file_path.suffix.lower()
    if suffix == ".docx":
        return _read_docx_text(file_path)

    if suffix == ".doc":
        # 尝试转换
        converted: Path | None = None
        try:
            converted = _convert_doc_to_docx_via_com(file_path)
            text = _read_docx_text(converted)
            return text
        finally:
            # 清理转换产物
            try:
                if converted and converted.exists():
                    converted.unlink()
            except Exception:
                # 清理失败不应影响主流程
                pass

    raise ValueError("仅支持 .doc 或 .docx 文件格式")


